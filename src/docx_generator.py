"""DOCX generation utilities."""

from pathlib import Path
from typing import Dict
from docx import Document


def replace_in_paragraph(paragraph, placeholders: Dict[str, str]) -> None:
    """Replace placeholders in a paragraph while preserving formatting as much as possible."""
    for key, value in placeholders.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(key, value)


def replace_in_table(table, placeholders: Dict[str, str]) -> None:
    """Replace placeholders in all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            replace_in_container(cell, placeholders)


def replace_in_container(container, placeholders: Dict[str, str]) -> None:
    """Replace placeholders in paragraphs and tables inside a container."""
    for paragraph in container.paragraphs:
        replace_in_paragraph(paragraph, placeholders)

    for table in container.tables:
        replace_in_table(table, placeholders)


def replace_placeholders(document: Document, placeholders: Dict[str, str]) -> None:
    """Replace placeholders in the document body, header, and footer."""
    replace_in_container(document, placeholders)

    for section in document.sections:
        replace_in_container(section.header, placeholders)
        replace_in_container(section.footer, placeholders)


def generate_docx(template_file: str | Path, output_file: str | Path, placeholders: Dict[str, str]) -> None:
    """Generate the DOCX file from the template."""
    document = Document(template_file)
    replace_placeholders(document, placeholders)
    document.save(output_file)
